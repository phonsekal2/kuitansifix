import streamlit as st
import pandas as pd
from docxtpl import DocxTemplate
from num2words import num2words
from datetime import datetime
from docx import Document
from docxcompose.composer import Composer
import numpy as np
def kalender_indo(value):
    a = (value).strftime("%d %B %Y")
    kal = a.replace("January", "Januari").replace("February", "Februari").replace("March", "Maret").replace("May", "Mei").replace("June", "Juni").replace("July", "Juli").replace("August", "Agustus").replace("October", "Oktober").replace("December", "Desember")
    return kal
def bulan_indo(value):
    a = (value)
    kal = a.replace("January", "Januari").replace("February", "Februari").replace("March", "Maret").replace("May", "Mei").replace("June", "Juni").replace("July", "Juli").replace("August", "Agustus").replace("October", "Oktober").replace("December", "Desember")
    return kal

def transform_to_rupiah_format(value):
    str_value = str(value)
    separate_decimal = str_value.split(".")
    after_decimal = separate_decimal[0]
    before_decimal = separate_decimal[1]
    reverse = after_decimal[::-1]
    temp_reverse_value = ""
    for index, val in enumerate(reverse):
        if (index + 1) % 3 == 0 and index + 1 != len(reverse):
            temp_reverse_value = temp_reverse_value + val + "."
        else:
            temp_reverse_value = temp_reverse_value + val
    temp_result = temp_reverse_value[::-1]
    return "Rp" + temp_result + ",-" 
    #+ "," + before_decimal ##buat koma di belakang

def rupiah_strip(value):
    a = str(transform_to_rupiah_format(float(value)))
    ubah = a.replace("Rp0,-", "Rp-")
    return ubah

st.set_page_config(page_title='Dede Saputra App', page_icon = ":coffee:", layout = 'centered', initial_sidebar_state = 'auto')

st.title("Aplikasi Kuitansi")
st.sidebar.title("Menu")
option = st.sidebar.selectbox("Pilih jenis kuitansi", ("PPK 01 (AKik Takjudin)", "PPK 02 (Syihabudin)"))
if option == 'PPK 01 (AKik Takjudin)':
    form = st.form("Upload file")
    with form:
        excelfile = st.file_uploader("Unggah file nominatif")
        submit = st.form_submit_button("Proses")
    if submit:
        df2 = pd.read_excel(excelfile, sheet_name="kerja")
        for r_idx, r_val in df2.iterrows():
            if (r_val['asal'] == 'Jakarta'):
                doctemp = f"pages/templates/jakartakuitansikegiatanppk.docx"
            elif (r_val['asal'] == 'Bogor'):
                doctemp = f"pages/templates/jakartakuitansikegiatanppk.docx"
            elif (r_val['asal'] == 'Banten'):
                doctemp = f"pages/templates/jakartakuitansikegiatanppk.docx"
            elif (r_val['asal'] == 'Bandung'):
                doctemp = f"pages/templates/jakartakuitansikegiatanppk.docx"
            elif (r_val['asal'] == 'Serang'):
                doctemp = f"pages/templates/jakartakuitansikegiatanppk.docx"
            else:
                doctemp = f"pages/templates/daerahkuitansikegiatanppk.docx"

            doc = DocxTemplate(doctemp)
            context = {
                'total' : rupiah_strip(r_val['tiket'] + r_val['taksi_jakarta'] + r_val['taksi_daerah'] + r_val['total_hari']),
                'output' : r_val['output'],
                'sub_output' : r_val['sub_output'],
                'akun' : r_val['akun'],
                'kegiatan' : r_val['kegiatan'],
                'layanan' : r_val['layanan'],
                'tempat' : r_val['tempat'],
                'tanggal' : bulan_indo(r_val['tanggal']),
                'nomor_st' : r_val['nomor_st'],
                'tanggal_st' : kalender_indo(r_val['tanggal_st']),
                'asal_tujuan' : r_val['asal'] + "-" + r_val['lokasi'],
                'terbilang' : num2words(int(r_val['tiket'] + r_val['taksi_jakarta'] + r_val['taksi_daerah'] + r_val['total_hari']), lang='id').title() + " Rupiah",
                #'terbilang' : rupiah_strip((r_val['tiket'] + r_val['taksi_jakarta'] + r_val['taksi_daerah'] + r_val['total_hari'])),
                'nama' : r_val['nama'],
                'nip' : r_val['nip'],
                'tiket' : rupiah_strip(r_val['tiket']),
                'taksi_jakarta' : rupiah_strip(r_val['taksi_jakarta']),
                'taksi_daerah' : rupiah_strip(r_val['taksi_daerah']),
                'hari' : str(r_val['hari']),
                'hari_rp' : rupiah_strip(float(r_val['total_hari'] / r_val['hari'])),
                'total_hari' : rupiah_strip(r_val['total_hari']),
                'lokasi' : r_val["lokasi"],
                'bulan' : kalender_indo(r_val['tanggal_akhir']),
                'jabatan' : r_val['jabatan'],
                'dpr_jakarta' : rupiah_strip(r_val['dpr_jakarta']),
                'dpr_daerah' : rupiah_strip(r_val['dpr_daerah']),
                'total_spd' : rupiah_strip(r_val['dpr_jakarta'] + r_val['dpr_daerah']),
                'mak' : r_val['mak'],                        
                }
            #st.write(context)
            doc.render(context)
            output_path = f"pages/OUTPUT/{context['nama']}.docx"
            doc.save(output_path)
        a = st.success("🎉 File kuitansi telah selesai dibuat, silahkan unduh")
        b = st.success("")
        with b:
            np_array = df2["nama"].to_numpy()
            #st.write(np_array)

            files2 = list("pages/OUTPUT/" + (np_array) + ".docx")
            composed = f"pages/gabung.docx"
            result = Document(files2[0])
            result.add_page_break()
            composer = Composer(result)
            for i in range(1, len(files2)):
                doc2 = Document(files2[i])
                if i != len(files2) -1:
                    doc2.add_page_break()
                composer.append(doc2)
            composer.save(composed)
            with open(composed, "rb") as file:
                st.success("🎉 File kuitansi telah selesai dibuat")
                st.download_button(
                    label = "⬇️ Download File",
                    data=file,
                    file_name="kuitansi.docx",
                    mime="application/octet-stream",
                    key="10000009"
                )
if option == 'PPK 02 (Syihabudin)':
    form = st.form("Upload file")
    with form:
        excelfile = st.file_uploader("Unggah file nominatif")
        submit = st.form_submit_button("Proses")
    if submit:
        df2 = pd.read_excel(excelfile, sheet_name="kerja")
        for r_idx, r_val in df2.iterrows():
            if (r_val['asal'] == 'Jakarta'):
                doctemp = f"pages/templates/jakartakuitansikegiatan.docx"
            elif (r_val['asal'] == 'Bogor'):
                doctemp = f"pages/templates/jakartakuitansikegiatan.docx"
            elif (r_val['asal'] == 'Banten'):
                doctemp = f"pages/templates/jakartakuitansikegiatan.docx"
            elif (r_val['asal'] == 'Bandung'):
                doctemp = f"pages/templates/jakartakuitansikegiatan.docx"
            elif (r_val['asal'] == 'Serang'):
                doctemp = f"pages/templates/jakartakuitansikegiatan.docx"
            else:
                doctemp = f"pages/templates/daerahkuitansikegiatan.docx"

            doc = DocxTemplate(doctemp)
            context = {
                'total' : rupiah_strip(r_val['tiket'] + r_val['taksi_jakarta'] + r_val['taksi_daerah'] + r_val['total_hari']),
                'output' : r_val['output'],
                'sub_output' : r_val['sub_output'],
                'akun' : r_val['akun'],
                'kegiatan' : r_val['kegiatan'],
                'layanan' : r_val['layanan'],
                'tempat' : r_val['tempat'],
                'tanggal' : bulan_indo(r_val['tanggal']),
                'nomor_st' : r_val['nomor_st'],
                'tanggal_st' : kalender_indo(r_val['tanggal_st']),
                'asal_tujuan' : r_val['asal'] + "-" + r_val['lokasi'],
                'terbilang' : num2words(int(r_val['tiket'] + r_val['taksi_jakarta'] + r_val['taksi_daerah'] + r_val['total_hari']), lang='id').title() + " Rupiah",
                #'terbilang' : rupiah_strip((r_val['tiket'] + r_val['taksi_jakarta'] + r_val['taksi_daerah'] + r_val['total_hari'])),
                'nama' : r_val['nama'],
                'nip' : r_val['nip'],
                'tiket' : rupiah_strip(r_val['tiket']),
                'taksi_jakarta' : rupiah_strip(r_val['taksi_jakarta']),
                'taksi_daerah' : rupiah_strip(r_val['taksi_daerah']),
                'hari' : str(r_val['hari']),
                'hari_rp' : rupiah_strip(float(r_val['total_hari'] / r_val['hari'])),
                'total_hari' : rupiah_strip(r_val['total_hari']),
                'lokasi' : r_val["lokasi"],
                'bulan' : kalender_indo(r_val['tanggal_akhir']),
                'jabatan' : r_val['jabatan'],
                'dpr_jakarta' : rupiah_strip(r_val['dpr_jakarta']),
                'dpr_daerah' : rupiah_strip(r_val['dpr_daerah']),
                'total_spd' : rupiah_strip(r_val['dpr_jakarta'] + r_val['dpr_daerah']),
                'mak' : r_val['mak'],                        
                }
            #st.write(context)
            doc.render(context)
            output_path = f"pages/OUTPUT/{context['nama']}.docx"
            doc.save(output_path)
        a = st.success("🎉 File kuitansi telah selesai dibuat, silahkan unduh")
        b = st.success("")
        with b:
            np_array = df2["nama"].to_numpy()
            #st.write(np_array)

            files2 = list("pages/OUTPUT/" + (np_array) + ".docx")
            composed = f"pages/gabung.docx"
            result = Document(files2[0])
            result.add_page_break()
            composer = Composer(result)
            for i in range(1, len(files2)):
                doc2 = Document(files2[i])
                if i != len(files2) -1:
                    doc2.add_page_break()
                composer.append(doc2)
            composer.save(composed)
            with open(composed, "rb") as file:
                st.success("🎉 File kuitansi telah selesai dibuat")
                st.download_button(
                    label = "⬇️ Download File",
                    data=file,
                    file_name="kuitansi.docx",
                    mime="application/octet-stream",
                    key="10000009"
                )
        
